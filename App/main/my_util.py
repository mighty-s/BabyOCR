__author__ = 'sight'
import cv2
import itertools
from docx import Document
from docx.shared import Inches
import pyttsx3

'''
    --------------------  Public Operations  ---------------------
    
    이 부분은 다른 파일에서 호출하는 함수 목록입니다.        
'''


def search_contours(contours, container, image):
    """
    Contours 들을 찾아서 contour 기반으로 사각형을 그리고
    리스트에 집어넣는 함수
    :param contours:    Contours 정보가 들어있는 리스트
    :param container:   Contour 정보를 토대로 사각형을 만들어서 좌표를 저장할 리스트
    :param image:      사각형을 표시할 이미지
    :return:           none
    """
    for i in range(len(contours)):
        cnt = contours[i]
        x, y, w, h = cv2.boundingRect(cnt)
        cv2.rectangle(image, (x, y), (x + w, y + h), (255, 0, 0), 1)  # 원본에 파랑 네모침 ( BGR 값 )
        container.append((x, y, w, h))


def letter_sort(boxes):
    """
    이미지에서 식별한 문자들을 정렬하는 함수
    문자의 위치 좌표기반으로 정렬하며,

    모든 문자들을 Y축 기준으로 정렬하고,
    이 중에서 같은 Y좌표 상에 있는 문자들만을 따로 뽑아내서 정렬후 리스트에 대입
    줄바꿈시 가운에 0 넣어줌
    [ ( 1번째 줄 ), 0 ,( 2번째 줄 ) .. ]
    
    이 리스트를 flat 시켜서 1차원 리스트로 리턴
    :param boxes:   박스 (x, y, w, h)를 엘리먼트로 가지고 있는 리스트
    :return:        (영어,한글 읽는 순서)로 정렬된 contour 배열
    """
    result = []                                    # 정렬된 contours들을 저장할 리스트
    boxes = sorted(boxes, key=lambda box: box[1])  # 박스들을 Y좌표 기준으로 정렬

    pre_line = 0                                   
    for i in range(len(boxes)):
        if i != 0:
            if _is_line_break(boxes[i], boxes[i-1]):    # 줄바꿈이 일어나면 시작점부터 해당 이미지까지 한줄로 처리
                line = _get_line(boxes, pre_line, i)
                result.append(line)
               # result.append(0)
                pre_line = i
        # 마지막줄 처리
        if i == len(boxes)-1:
            line = _get_line(boxes, pre_line, i+1)
            result.append(line)

    return list(itertools.chain.from_iterable(result))         # 2차원 배열 flat 시켜서 리턴


def write_in_docx(text):
    """
    읽어들인 문자를 워드 형식으로 저장해주는 함수
    :param text:  이미지에서 추출한 문자열
    :return:      none
    """
    document = Document()
    document.add_heading('변환결과 ', 0)
    p = document.add_paragraph(text)
    document.save('./output/result.docx')


def text_to_speech(text):
    """
    읽어들인 문자를 음성으로 재생하주는 함수
    :param text:    이미지에서 추출한 문자열
    :return:        none
    """
    engine = pyttsx3.init()
    voices = engine.getProperty('voices')
    engine.say('다음은 문서에서 추출된 음성의 내용입니다. ')
    engine.say(text)
    engine.runAndWait()


'''
    --------------------  Private Operations  ---------------------
    
    이 부분 부터 정의된 함수는 이 파일 밖에서 쓰이지 않습니다.
'''


def _is_line_break(contour1, contour2):
    """
    이미지 상에서 글씨 줄바꿈이 일어났는지 판별하는 함수
    한글, 영어등 왼쪽에서 오른쪽으로 읽는 문자들에게만 유효하다.
    해당 contours 의 Y 좌표의 차이가 40픽셀 이상일 경우 줄바꿈 처리

    :param contour1: 비교하려는 왼족 글자 contour
    :param contour2: 비교하려는 오른쪽 글자 contour
    :return: 줄바꿈이 일어났는지 여부
    """
    return contour1[1] - contour2[1] >= 20


def _get_line(original, start, end):
    """
    
    :param original:  모든 contour들을 가지고 있는 원본 배열 ( Y 좌표 기준으로 정렬된 상태 )
    :param start:     자르기 시작할 인덱스 번호 ( inclusive )
    :param end:       끝까지
    :return: 
    """
    line = original[start: end]
    line.sort(key=lambda char: char[0])    # 같은 줄에 있는 contours x 좌표 기준으로 정렬
    return line

